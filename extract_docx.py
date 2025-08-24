import docx
import sys

def extract_text_from_docx(file_path):
    doc = docx.Document(file_path)
    full_text = []
    
    # Extract text from paragraphs
    for para in doc.paragraphs:
        full_text.append(para.text)
    
    # Extract text from tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    full_text.append(paragraph.text)
    
    return '\n'.join(full_text)

if __name__ == "__main__":
    if len(sys.argv) != 3:
        print("Usage: python extract_docx.py input_file.docx output_file.txt")
        sys.exit(1)
    
    input_file = sys.argv[1]
    output_file = sys.argv[2]
    
    try:
        text = extract_text_from_docx(input_file)
        with open(output_file, 'w', encoding='utf-8') as f:
            f.write(text)
        print(f"Text successfully extracted from {input_file} to {output_file}")
    except Exception as e:
        print(f"Error: {e}")
        sys.exit(1)